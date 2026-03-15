import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_documento_word(titulo_tesis, contenido_redactado, bibliografia=""):
    doc = Document()
    
    # Título principal
    titulo = doc.add_heading(titulo_tesis, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Añadir los capítulos ordenados
    for n in sorted(contenido_redactado.keys(), key=lambda x: int(x)):
        doc.add_heading(f"Capítulo {n}", level=1)
        parrafo = doc.add_paragraph(contenido_redactado[n])
        parrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    # Añadir Bibliografía si existe
    if bibliografia:
        doc.add_page_break()
        doc.add_heading("Bibliografía", level=1)
        parrafo_bib = doc.add_paragraph(bibliografia)
        parrafo_bib.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
    # Guardar en memoria para descarga en Streamlit
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
