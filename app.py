import streamlit as st
from supabase import create_client, Client
import google.generativeai as genai
import json
import io 
from docx import Document 
import re 

# --- 1. CONFIGURACIÓN Y DICCIONARIO DE IDIOMAS (ACTUALIZADO) ---

TRANSLATIONS = {
    "Español": {
        "page_title": "Buscador de 訓詁",
        "main_title": "🏯 Buscador de 訓詁 (Xùngǔ)",
        "desc": "Consulta múltiples fuentes clásicas y genera análisis con IA.",
        "db_select": "Bases de datos a consultar",
        "input_req": "Prompt / Petición (Contexto para el análisis)",
        "input_req_placeholder": "Ej: Analiza el concepto de virtud...",
        "input_keywords": "Palabras Clave para búsqueda (separadas por comas)",
        "input_keywords_placeholder": "Ej: Confucio, virtud, Ren",
        "resp_lang": "Idioma de la respuesta (IA)",
        "btn_analyze": "🔍 Buscar y Analizar",
        "warn_input": "Por favor, introduce el prompt, las palabras clave y selecciona una base de datos.",
        "searching": "Buscando en Supabase...",
        "error_not_found": "No se encontraron datos con esas palabras clave.",
        "success_found": "¡Contexto recuperado correctamente!",
        "analyzing": "Consultando a Gemini 2.0 Flash...",
        "result_title": "📜 Resultado del Análisis",
        "source_title": "Ver datos JSON enviados a la IA",
        "btn_download_word": "📥 Descargar Análisis en Word",
        "filename_prefix": "Analisis_Xungu",
        "sidebar_lang": "Idioma de la Interfaz / 介面語言"
    },
    "Traditional Chinese": {
        "page_title": "訓詁搜尋器",
        "main_title": "🏯 訓詁搜尋器 (Xùngǔ)",
        "desc": "查詢多種經典文獻並透過 AI 生成分析。",
        "db_select": "要查詢的資料庫",
        "input_req": "提示詞 / 具體要求",
        "input_req_placeholder": "例如：分析這些文本中的美德概念...",
        "input_keywords": "搜尋關鍵字 (用逗號隔開)",
        "input_keywords_placeholder": "例如：孔子, 美德, 仁",
        "resp_lang": "回覆語言 (AI)",
        "btn_analyze": "🔍 搜尋並分析",
        "warn_input": "請輸入提示詞、關鍵字並選擇資料庫。",
        "searching": "正在搜尋資料庫...",
        "error_not_found": "找不到相符的資料。",
        "success_found": "成功載入上下文！",
        "analyzing": "正在諮詢 Gemini 2.0 Flash...",
        "result_title": "📜 分析結果",
        "source_title": "查看發送給 AI 的 JSON 資料",
        "btn_download_word": "📥 下載 Word 分析報告",
        "filename_prefix": "Xungu_Analysis",
        "sidebar_lang": "Interface Language / 介面語言"
    },
    "English": {
        "page_title": "Xùngǔ Searcher",
        "main_title": "🏯 Xùngǔ Searcher (Exegesis)",
        "desc": "Search multiple classical sources and generate AI analysis.",
        "db_select": "Databases to query",
        "input_req": "Prompt / Request",
        "input_req_placeholder": "E.g.: Analyze the concept of virtue...",
        "input_keywords": "Search Keywords (separated by commas)",
        "input_keywords_placeholder": "E.g.: Confucius, virtue, Ren",
        "resp_lang": "Response Language (AI)",
        "btn_analyze": "🔍 Search & Analyze",
        "warn_input": "Please enter prompt, keywords, and select a database.",
        "searching": "Searching Supabase...",
        "error_not_found": "No data found with those keywords.",
        "success_found": "Context loaded successfully!",
        "analyzing": "Consulting Gemini 2.0 Flash...",
        "result_title": "📜 Analysis Result",
        "source_title": "View JSON data sent to AI",
        "btn_download_word": "📥 Download Analysis as Word",
        "filename_prefix": "Xungu_Analysis",
        "sidebar_lang": "Interface Language / 介面語言"
    }
}

st.set_page_config(page_title="Sinología AI", layout="centered")

# --- 2. SELECTOR DE IDIOMA ---
idiomas_disponibles = ["Español", "Traditional Chinese", "English"]
lang_sel = st.sidebar.selectbox("Language / Idioma / 語言", idiomas_disponibles)
T = TRANSLATIONS[lang_sel]

TABLAS_DISPONIBLES = [
    "Glosas de 鬼谷子",
    "Analectas de Confucio",
    "Fuentes secundarias",
    "JSON de investigación",
    "Mencio",
    "Xunzi",
    "戰國策"
]

# --- 3. CONEXIÓN A SUPABASE Y GEMINI ---
try:
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    genai.configure(api_key=GOOGLE_API_KEY)
except Exception as e:
    st.error(f"Error config: {e}")
    st.stop()

# --- 4. FUNCIÓN HELPER PARA WORD ---
def crear_word(titulo, subtitulo, contenido):
    doc = Document()
    doc.add_heading(titulo, 0)
    doc.add_heading(subtitulo, level=1)
    doc.add_paragraph(contenido)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 5. INTERFAZ DE USUARIO ---
st.title(T["main_title"])
st.markdown(T["desc"])

with st.form("research_form"):
    tablas_seleccionadas = st.multiselect(
        T["db_select"], 
        options=TABLAS_DISPONIBLES, 
        default=[] 
    )

    # NUEVO: Campo para palabras clave manuales
    keywords_manuales = st.text_input(
        T["input_keywords"], 
        placeholder=T["input_keywords_placeholder"]
    )

    peticion_concreta = st.text_area(
        T["input_req"], 
        placeholder=T["input_req_placeholder"], 
        height=100
    )

    idioma_salida = st.selectbox(
        T["resp_lang"],
        options=["Español", "English", "中文 (Traditional Chinese)", "Français"]
    )

    submitted = st.form_submit_button(T["btn_analyze"])

# --- 6. LÓGICA DEL BACKEND ---
if submitted:
    # Validamos que existan palabras clave además del prompt
    if not peticion_concreta or not tablas_seleccionadas or not keywords_manuales:
        st.warning(T["warn_input"])
    else:
        with st.spinner(T["searching"]):
            contexto_encontrado = []
            
            # Procesamos las palabras clave manuales
            # Las separamos por comas y limpiamos espacios
            lista_keywords = [k.strip() for k in keywords_manuales.split(",") if k.strip()]
            
            try:
                for tabla in tablas_seleccionadas:
                    query = supabase.table(tabla).select("*")
                    
                    # Aplicamos el filtro de palabras clave si existen
                    if lista_keywords:
                        # Buscamos en la columna "Palabras Clave" (asegúrate que se llame así en todas tus tablas)
                        condiciones_or = ",".join([f'"Palabras Clave".ilike.%{kw}%' for kw in lista_keywords])
                        query = query.or_(condiciones_or)
                    
                    response = query.execute()
                    
                    if response.data:
                        contexto_encontrado.append({
                            "Fuente/Tabla": tabla,
                            "Datos": response.data
                        })
                
                if not contexto_encontrado:
                    st.error(T["error_not_found"])
                    st.stop()
                
                st.success(f"{T['success_found']} ({len(contexto_encontrado)} fuentes con resultados)")
                
            except Exception as e:
                st.error(f"Error Supabase: {e}")
                st.stop()

        # --- 7. GENERACIÓN CON GEMINI ---
        with st.spinner(T["analyzing"]):
            try:
                contexto_texto = json.dumps(contexto_encontrado, indent=2, ensure_ascii=False)

                prompt_final = f"""
                Role: You are an Expert Sinologist specializing in classical Chinese texts and 'Xungu' (Exegesis).
                USER PROMPT: "{peticion_concreta}"
                RETRIEVED CONTEXT:
                ```json
                {contexto_texto}
                ```
                INSTRUCTIONS:
                1. Answer based PRIMARILY on the provided JSON.
                2. Use EXPLICIT CITATIONS (author, year) in the text.
                3. Add a Bibliography at the end using the citation fields in the JSON.
                4. RESPONSE LANGUAGE: {idioma_salida}.
                """

                model = genai.GenerativeModel('gemini-2.0-flash') 
                response_ai = model.generate_content(prompt_final)
                
                st.markdown(f"### {T['result_title']}")
                st.write(response_ai.text)
                
                # BOTÓN DE DESCARGA
                word_file = crear_word(T["main_title"], "Análisis de Fuentes", response_ai.text)
                st.download_button(
                    label=T["btn_download_word"],
                    data=word_file,
                    file_name=f"{T['filename_prefix']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                with st.expander(T["source_title"]):
                    st.json(contexto_encontrado)

            except Exception as e:
                st.error(f"Error Gemini: {e}")
